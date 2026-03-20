from gettext import install

import pandas as pd
import numpy as np
from sklearn.preprocessing import LabelEncoder, StandardScaler

from docx import Document
from docx.shared import Pt

print("=" * 60)
print("STEP 1: LOADING DATASET")
print("=" * 60)

df = pd.read_excel(r'C:\Users\kotha\OneDrive\Desktop\ASC WIL\my_data.xlsx')

print(f"shape: {df.shape[0]} rows, {df.shape[1]} columns")
print(f"\nColumns: {list(df.columns)}")

print("\n" + "=" * 60)
print("STEP 2: INITIAL DATA INSPECTION")
print("=" * 60)

print("\nData Types:")
print(df.dtypes)

print("\nFirst 10 rows:")
print(df.head())

print("\n" + "=" * 60)
print("STEP 3: CHECK AND DROP DUPLICATES")
print("=" * 60)

# ── How many duplicates exist ─────────────────────────────
total_rows_before = len(df)
duplicate_count   = df.duplicated().sum()
duplicate_pct     = (duplicate_count / total_rows_before) * 100

print(f"\nTotal rows before : {total_rows_before}")
print(f"Duplicate rows    : {duplicate_count} ({duplicate_pct:.2f}% of dataset)")

# ── Show a sample of the actual duplicate rows ────────────
if duplicate_count > 0:
    print("\nSample duplicate rows (showing first 4):")
    sample_dupes = df[df.duplicated(keep=False)].sort_values(
                   list(df.columns)).head(4)
    print(sample_dupes.to_string())

df = df.drop_duplicates()
total_rows_after = len(df)

print(f"\nTotal rows after dropping : {total_rows_after}")
print(f"Rows removed: {total_rows_before - total_rows_after}")
print(f"Duplicates dropped successfully!")

print("\n" + "=" * 45)
print("STEP 3: MISSING VALUES")
print("=" * 45)

missing = df.isnull().sum()
print("\nMissing values per column:")
print(missing)

print("\n" + "=" * 45)
print("STEP 4: RENAME COLUMNS")
print("=" * 45)

df = df.rename(columns={
    'gender': 'Gender',
    'tenure': 'Tenure in Months',
    'SeniorCitizen': ('Senior Citizen'),
    'PhoneService': ('Phone Service'),
    'MultipleLines': ('Multiple Lines'),
    'InternetService': ('Internet Service'),
    'MonthlyCharges': ('Monthly Charges in $'),
    'Dependents': ('Dependants'),
})

print("\nColumn names after renaming:")
print(df.columns.tolist())

print("\n" + "=" * 50)
print("STEP 5A: LABEL ENCODING (Binary Columns)")
print("=" * 50)

le = LabelEncoder()

# Gender: Female=0, Male=1
df['Gender'] = le.fit_transform(df['Gender'])
print("\nGender        → Female=0, Male=1")

# Dependents: No=0, Yes=1
df['Dependants'] = le.fit_transform(df['Dependants'])
print("Dependants    → No=0, Yes=1")

# PhoneService: No=0, Yes=1
df['Phone Service'] = le.fit_transform(df['Phone Service'])
print("Phone Service  → No=0, Yes=1")

# MultipleLines: No=0, Yes=1
df['Multiple Lines'] = le.fit_transform(df['Multiple Lines'])
print("Multiple Lines → No=0, Yes=1")

# Churn: No=0, Yes=1
df['Churn'] = le.fit_transform(df['Churn'])
print("Churn         → No=0, Yes=1")

# ============================================================
# STEP 5B: ONE-HOT ENCODING
# For Contract column (3 categories)
# ============================================================
print("\n" + "=" * 50)
print("STEP 5B: ONE-HOT ENCODING (Contract Column, Internet Services Coloumn)")
print("=" * 50)

print("\nOriginal Contract unique values:")
print(df['Contract'].unique())

df = pd.get_dummies(df, columns=['Contract','Internet Service'], prefix=['Contract','Internet Service'])

# Captures both 'Contract_...' AND 'Internet Service_...'
ohe_cols = [col for col in df.columns
            if col.startswith('Contract') or col.startswith('Internet Service')]
df[ohe_cols] = df[ohe_cols].astype(int)

print("\nNew Contract columns created:")
contract_cols = [col for col in df.columns if 'Contract' in col]
print(contract_cols)
print("\nSample (first 5 rows):")
print(df[contract_cols].head())

internet_cols = [col for col in df.columns if col.startswith('Internet Service')]
print("\nNew Internet Service columns created:")
print(internet_cols)
print("\nSample Internet Service (first 5 rows):")
print(df[internet_cols].head())


# ============================================================
# FINAL RESULT
# ============================================================
print("\n" + "=" * 50)
print("FINAL ENCODED DATASET")
print("=" * 50)

print("\nShape:", df.shape)
print("\nAll Columns:")
print(list(df.columns))
print("\nData Types:")
print(df.dtypes)
print("\nFirst 5 rows:")
print(df.head())

# ============================================================
# Step 6: FEATURE ENGINEERING
# ============================================================

df_model = df.copy()
df_model['Tenure_Group'] = pd.cut(
    df_model['Tenure in Months'],
    bins=[0,12,24,48,72],
    labels=['New','Short','Medium','Loyal']
)

# Fixed Charge Category to match
df_model['Charge_Category'] = pd.cut(
    df_model['Monthly Charges in $'],
    bins=[0, 40, 70, 90, 120],
    labels=['Low','Medium','High','Very High']
)

# Fixed High Risk Flag - now checks for string labels
df_model['High_Risk_Flag'] = (
    (df_model['Tenure_Group'] == 'New') &
    (df_model['Charge_Category'].isin(['High', 'Very High']))
).astype(int)

print(f"✅ High Risk Flag created ({df_model['High_Risk_Flag'].sum()} customers flagged)")

# Fixed Loyalty Score - convert strings to numbers
print("\nCreating Feature 4: Loyalty_Score")
tenure_map = {'New': 0, 'Short': 1, 'Medium': 2, 'Loyal': 3}

# Map and fill any NaN with 0 before converting to int
df_model['Loyalty_Score'] = (
    df_model['Tenure_Group'].map(tenure_map).fillna(0).astype(int) +  # ← Add .fillna(0)
    df_model['Contract_Two year'] * 2 +
    df_model['Contract_One year'] * 1
)

print(f"✅ Loyalty Score created (range: {df_model['Loyalty_Score'].min()}-{df_model['Loyalty_Score'].max()})")
print(f"\nTotal features after engineering: {df_model.shape[1]}")


# ============================================================
# STEP 7: ENCODE STRING CATEGORIES FOR ML
# ============================================================
print("\n" + "=" * 70)
print("STEP 3B: ENCODE STRING CATEGORIES FOR ML")
print("=" * 70)

# For ML models, we need numeric values
# Create encoded versions of the categorical features
print("\nEncoding Tenure_Group and Charge_Category...")

from sklearn.preprocessing import LabelEncoder
le_tenure = LabelEncoder()
le_charge = LabelEncoder()

# Handle potential NaN values before encoding
print(f"\nChecking for NaN before encoding:")
print(f"  Tenure_Group NaN: {df_model['Tenure_Group'].isna().sum()}")
print(f"  Charge_Category NaN: {df_model['Charge_Category'].isna().sum()}")

# Convert categorical to string to avoid issues, fill NaN if any
df_model['Tenure_Group_encoded'] = le_tenure.fit_transform(
    df_model['Tenure_Group'].astype(str)
)
df_model['Charge_Category_encoded'] = le_charge.fit_transform(
    df_model['Charge_Category'].astype(str)
)

print(f"\nTenure_Group encoding:")
for orig, encoded in zip(le_tenure.classes_, le_tenure.transform(le_tenure.classes_)):
    print(f"  {orig:<10} → {encoded}")

print(f"\nCharge_Category encoding:")
for orig, encoded in zip(le_charge.classes_, le_charge.transform(le_charge.classes_)):
    print(f"  {orig:<12} → {encoded}")

print("\n✅ All features ready for ML modeling!")
print(f"\nColumns available for modeling: {df_model.shape[1]}")

# ============================================================
# STEP 8: SEPARATE FEATURES AND TARGET
# ============================================================
print("\n" + "=" * 70)
print("STEP 4: SEPARATE FEATURES AND TARGET")
print("=" * 70)

# Drop the original string columns (keep only encoded versions)
# Also drop the original Tenure and Charges columns used to create groups
X = df_model.drop(columns=[
    'Churn',  # Target variable
    'Tenure_Group',  # String version (we have encoded)
    'Charge_Category',  # String version (we have encoded)
])

y = df_model['Churn']

print(f"\nFeatures (X): {X.shape}")
print(f"Target (y)  : {y.shape}")

print(f"\nTarget distribution:")
print(f"  No churn (0): {(y==0).sum()} ({(y==0).mean()*100:.1f}%)")
print(f"  Churn (1)   : {(y==1).sum()} ({(y==1).mean()*100:.1f}%)")

print(f"\nAll feature columns:")
for i, col in enumerate(X.columns, 1):
    print(f"  {i:2d}. {col}")

# ============================================================
# STEP 9: TRAIN/TEST SPLIT
# ============================================================
print("\n" + "=" * 70)
print("STEP 5: TRAIN/TEST SPLIT")
print("=" * 70)

from sklearn.model_selection import train_test_split

X_train, X_test, y_train, y_test = train_test_split(
    X, y,
    test_size=0.2,       # 80% train, 20% test
    random_state=42,     # reproducibility
    stratify=y           # maintain churn ratio in both sets
)

print(f"\nDataset split:")
print(f"  Total samples     : {len(X)}")
print(f"  Training samples  : {len(X_train)} ({len(X_train)/len(X)*100:.1f}%)")
print(f"  Testing samples   : {len(X_test)} ({len(X_test)/len(X)*100:.1f}%)")

print(f"\nChurn distribution in training set:")
print(f"  No churn (0): {(y_train==0).sum()} ({(y_train==0).mean()*100:.1f}%)")
print(f"  Churn (1)   : {(y_train==1).sum()} ({(y_train==1).mean()*100:.1f}%)")

print(f"\nChurn distribution in testing set:")
print(f"  No churn (0): {(y_test==0).sum()} ({(y_test==0).mean()*100:.1f}%)")
print(f"  Churn (1)   : {(y_test==1).sum()} ({(y_test==1).mean()*100:.1f}%)")

print("\n✅ stratify=y ensures both sets have same churn ratio")

# ============================================================
# STEP 10: FEATURE SCALING
# ============================================================
print("\n" + "=" * 70)
print("STEP 6: FEATURE SCALING")
print("=" * 70)

# Identify columns to scale
# Wide numeric ranges need scaling, binary (0/1) don't
cols_to_scale = ['Tenure in Months', 'Monthly Charges in $', 'Loyalty_Score']

# Check if columns exist
cols_to_scale = [col for col in cols_to_scale if col in X_train.columns]

print(f"\nFeatures to scale ({len(cols_to_scale)}):")
for col in cols_to_scale:
    print(f"  ✅ {col:<25} range: {X_train[col].min():.0f} to {X_train[col].max():.0f}")

print(f"\nFeatures NOT scaled ({X_train.shape[1] - len(cols_to_scale)}):")
print("  ⏭️  All binary features (0/1) — no scaling needed")

print("\nBEFORE scaling (training set):")
print(X_train[cols_to_scale].describe().round(2))

# Fit scaler on TRAINING data only
from sklearn.preprocessing import StandardScaler
scaler = StandardScaler()
scaler.fit(X_train[cols_to_scale])

print(f"\nScaler fitted on TRAINING data only:")
print(f"  Mean values: {scaler.mean_.round(2)}")
print(f"  Std values : {scaler.scale_.round(2)}")

# Transform both sets
X_train_scaled = X_train.copy()
X_test_scaled = X_test.copy()

X_train_scaled[cols_to_scale] = scaler.transform(X_train[cols_to_scale])
X_test_scaled[cols_to_scale] = scaler.transform(X_test[cols_to_scale])

print(f"\nAFTER scaling (training set):")
print(X_train_scaled[cols_to_scale].describe().round(2))

print("\n  IMPORTANT:")
print("  • Scaler fitted on TRAINING data only")
print("  • Same scaler applied to BOTH train and test")
print("  • This prevents data leakage")


# ============================================================
# STEP 7: SAVE FILES
# ============================================================
print("\n" + "=" * 70)
print("STEP 7: SAVE TRAINING AND TESTING SETS")
print("=" * 70)

# Combine X and y for saving
train_set = X_train_scaled.copy()
train_set['Churn'] = y_train.values

test_set = X_test_scaled.copy()
test_set['Churn'] = y_test.values

# Save scaled versions
train_set.to_csv('churn_train_scaled.csv', index=False)
test_set.to_csv('churn_test_scaled.csv', index=False)

print(f"\n✅ Saved: churn_train_scaled.csv ({train_set.shape})")
print(f"✅ Saved: churn_test_scaled.csv ({test_set.shape})")

# ============================================================
# STEP 8A: TRAIN & TEST DOCUMENTATION (WORD FILE)
# ============================================================

doc = Document()

doc.add_heading('Training and Testing Sets Documentation', level=1)

doc.add_heading('Dataset Information', level=2)
doc.add_paragraph(f"Total samples: {len(X)}")
doc.add_paragraph(f"Total features: {X.shape[1]}")
doc.add_paragraph("Target variable: Churn (0 = No, 1 = Yes)")

doc.add_heading('Split Configuration', level=2)
doc.add_paragraph("Split ratio: 80% Training / 20% Testing")
doc.add_paragraph("Random state: 42")
doc.add_paragraph("Stratified split applied to preserve churn distribution")

doc.add_heading('Training Set', level=2)
doc.add_paragraph(f"File: churn_train_scaled.csv")
doc.add_paragraph(f"Size: {len(train_set)} rows × {train_set.shape[1]} columns")
doc.add_paragraph(
    f"No churn: {(y_train==0).sum()} ({(y_train==0).mean()*100:.1f}%)"
)
doc.add_paragraph(
    f"Churn: {(y_train==1).sum()} ({(y_train==1).mean()*100:.1f}%)"
)

doc.add_heading('Testing Set', level=2)
doc.add_paragraph(f"File: churn_test_scaled.csv")
doc.add_paragraph(f"Size: {len(test_set)} rows × {test_set.shape[1]} columns")
doc.add_paragraph(
    f"No churn: {(y_test==0).sum()} ({(y_test==0).mean()*100:.1f}%)"
)
doc.add_paragraph(
    f"Churn: {(y_test==1).sum()} ({(y_test==1).mean()*100:.1f}%)"
)

doc.add_heading('Feature Engineering Applied', level=2)
doc.add_paragraph("• Tenure_Group_encoded")
doc.add_paragraph("• Charge_Category_encoded")
doc.add_paragraph("• High_Risk_Flag")
doc.add_paragraph("• Loyalty_Score")

doc.add_heading('Generated Files', level=2)
doc.add_paragraph("1. churn_train_scaled.csv")
doc.add_paragraph("2. churn_test_scaled.csv")
doc.add_paragraph("3. churn_train_unscaled.csv")
doc.add_paragraph("4. churn_test_unscaled.csv")

doc.save("Training_Testing_Documentation.docx")

print("✅ Training & Testing documentation saved.")

# ============================================================
# STEP 8B: SCALING TECHNIQUES DOCUMENTATION
# ============================================================

scaling_doc = Document()

scaling_doc.add_heading('Scaling Techniques Documentation', level=1)

scaling_doc.add_heading('Why Scaling Was Required', level=2)
scaling_doc.add_paragraph(
"Clustering and predictive models are distance-based algorithms. "
"Features with larger numeric ranges dominate model learning. "
"Scaling ensures equal contribution of all numerical variables."
)

scaling_doc.add_heading('Scaling Method Used', level=2)
scaling_doc.add_paragraph("StandardScaler (Mean = 0, Standard Deviation = 1)")

scaling_doc.add_heading('Features Scaled', level=2)
for col in cols_to_scale:
    scaling_doc.add_paragraph(col)

scaling_doc.add_heading('Scaling Procedure', level=2)
scaling_doc.add_paragraph(
"Scaler was fitted ONLY on the training dataset to prevent data leakage."
)

scaling_doc.add_heading('Code Snippet Used', level=2)

code_text = f"""
from sklearn.preprocessing import StandardScaler

scaler = StandardScaler()

X_train[{cols_to_scale}] = scaler.fit_transform(X_train[{cols_to_scale}])
X_test[{cols_to_scale}] = scaler.transform(X_test[{cols_to_scale}])
"""

scaling_doc.add_paragraph(code_text)

scaling_doc.add_heading('Outcome', level=2)
scaling_doc.add_paragraph(
"All scaled features now have mean ≈ 0 and standard deviation ≈ 1."
)

scaling_doc.save("Scaling_Techniques_Documentation.docx")

print("✅ Scaling documentation saved.")
